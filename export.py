"""
Exportación de contratos a Word (.docx) y PDF.
"""

import io
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


def _parse_lines(contract_text: str) -> list[dict]:
    """
    Analiza el texto del contrato y clasifica cada línea por tipo:
    - title: encabezado principal (CONTRATO DE...)
    - heading: nombre de sección (DECLARACIONES, CLÁUSULAS, PRIMERA., etc.)
    - body: texto normal
    - blank: línea vacía
    """
    lines = contract_text.split("\n")
    parsed = []

    heading_patterns = [
        r"^D\s*E\s*C\s*L\s*A\s*R\s*A\s*C\s*I\s*O\s*N\s*E\s*S",
        r"^C\s*L\s*[ÁA]\s*U\s*S\s*U\s*L\s*A\s*S",
        r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA|"
        r"DÉCIM[AO]|UNDÉCIM[AO]|DUODÉCIM[AO]|VIGÉSIM[AO])\b",
        r"^[IVX]+\.\s+Declara",
        r"^III\.\s+Declaran",
    ]

    for line in lines:
        stripped = line.strip()

        if not stripped:
            parsed.append({"type": "blank", "text": ""})
            continue

        # Título principal
        if stripped.upper().startswith("CONTRATO DE") and len(stripped) < 300:
            parsed.append({"type": "title", "text": stripped})
            continue

        # Headings
        is_heading = False
        for pattern in heading_patterns:
            if re.match(pattern, stripped, re.IGNORECASE):
                is_heading = True
                break

        if is_heading:
            parsed.append({"type": "heading", "text": stripped})
        else:
            parsed.append({"type": "body", "text": stripped})

    return parsed


def contract_to_docx(contract_text: str) -> bytes:
    """Convierte el texto del contrato a un documento Word (.docx)."""
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Arial"
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.space_after = Pt(4)
    style_normal.paragraph_format.line_spacing = 1.15

    parsed = _parse_lines(contract_text)

    for item in parsed:
        if item["type"] == "blank":
            doc.add_paragraph("")
            continue

        if item["type"] == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["text"])
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = "Arial"
            continue

        if item["type"] == "heading":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item["text"])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            continue

        # body
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item["text"])
        run.font.size = Pt(11)
        run.font.name = "Arial"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


class _ContractPDF(FPDF):
    """PDF con header/footer para contratos legales."""

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Página {self.page_no()}/{{nb}}", align="C")


def contract_to_pdf(contract_text: str) -> bytes:
    """Convierte el texto del contrato a PDF."""
    pdf = _ContractPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Márgenes
    pdf.set_left_margin(30)
    pdf.set_right_margin(25)

    parsed = _parse_lines(contract_text)

    for item in parsed:
        if item["type"] == "blank":
            pdf.ln(4)
            continue

        if item["type"] == "title":
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 7, item["text"], align="C")
            pdf.ln(3)
            continue

        if item["type"] == "heading":
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, item["text"], align="J")
            pdf.ln(2)
            continue

        # body
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, item["text"], align="J")
        pdf.ln(1)

    return bytes(pdf.output())
